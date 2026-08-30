"""
Script to create pilot test material DOCX file.

Generates: "T 01 PRINCÍPIOS ELEMENTARES DA MATÉRIA3.docx"
Content: Temporada 01 with 6 episodes and 15 questions + alternatives + answer key.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_pilot_material():
    """Create the pilot DOCX document."""
    doc = Document()
    doc.core_properties.title = "Temporada 01 — Princípios Elementares da Matéria"
    doc.core_properties.author = "AGENTE IA EDU"

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("T 01 — PRINCÍPIOS ELEMENTARES DA MATÉRIA")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Temporada 01 — Princípios Elementares da Matéria")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Season intro
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Nesta temporada, exploraremos os conceitos fundamentais da matéria, "
        "desde sua estrutura básica até as propriedades que definem tudo que nos rodeia."
    )
    intro_run.font.italic = True

    doc.add_paragraph()

    # Episodes and questions
    episodes = [
        {
            "number": 1,
            "title": "Estrutura da Matéria",
            "description": "Descobrindo os blocos de construção do universo",
        },
        {
            "number": 2,
            "title": "Átomos e Moléculas",
            "description": "A combinação de elementos fundamentais",
        },
        {
            "number": 3,
            "title": "Estados da Matéria",
            "description": "Sólido, líquido e gasoso",
        },
        {
            "number": 4,
            "title": "Ligações Químicas",
            "description": "Como os átomos se unem",
        },
        {
            "number": 5,
            "title": "Propriedades e Transformações",
            "description": "Mudanças físicas e químicas",
        },
        {
            "number": 6,
            "title": "A Tabela Periódica",
            "description": "Organizando os elementos",
        },
    ]

    question_bank = [
        {
            "text": "Qual é a unidade fundamental de toda a matéria?",
            "alternatives": ["A) Molécula", "B) Átomo", "C) Elétron", "D) Próton", "E) Nêutron"],
            "answer": "B",
            "explanation": "O átomo é a unidade fundamental de toda a matéria, sendo a menor parte de um elemento que mantém suas propriedades.",
        },
        {
            "text": "Um elétron tem carga negativa. Qual é a carga do próton?",
            "alternatives": ["A) Negativa", "B) Positiva", "C) Neutra", "D) Dupla", "E) Variável"],
            "answer": "B",
            "explanation": "O próton possui carga positiva, sendo responsável pelo núcleo atômico.",
        },
        {
            "text": "Quantos prótons possui um átomo de carbono?",
            "alternatives": ["A) 4", "B) 6", "C) 8", "D) 12", "E) 14"],
            "answer": "B",
            "explanation": "O carbono possui número atômico 6, portanto 6 prótons.",
        },
        {
            "text": "Qual é o estado da matéria que tem forma e volume definidos?",
            "alternatives": ["A) Gás", "B) Líquido", "C) Sólido", "D) Plasma", "E) Todos possuem"],
            "answer": "C",
            "explanation": "Os sólidos têm forma e volume bem definidos devido à forte coesão molecular.",
        },
        {
            "text": "Em qual estado da matéria as moléculas estão mais afastadas?",
            "alternatives": ["A) Sólido", "B) Líquido", "C) Gás", "D) Plasma", "E) Condensado"],
            "answer": "C",
            "explanation": "No estado gasoso, as moléculas estão muito afastadas e em movimento desordenado.",
        },
        {
            "text": "O que acontece quando congelamos água?",
            "alternatives": ["A) Mudança química", "B) Mudança física", "C) Transformação", "D) Reação", "E) Nenhuma"],
            "answer": "B",
            "explanation": "Congelar é uma mudança física, pois a água permanece sendo H₂O.",
        },
        {
            "text": "Qual é o metal mais abundante na Terra?",
            "alternatives": ["A) Ouro", "B) Prata", "C) Ferro", "D) Alumínio", "E) Cobre"],
            "answer": "D",
            "explanation": "O alumínio é o metal mais abundante na crosta terrestre.",
        },
        {
            "text": "Qual ligação ocorre entre H e O na água?",
            "alternatives": ["A) Iônica", "B) Covalente", "C) Metálica", "D) Eletrostática", "E) Magnética"],
            "answer": "B",
            "explanation": "A água é formada por ligações covalentes entre hidrogênio e oxigênio.",
        },
        {
            "text": "Qual é a temperatura de ebulição da água ao nível do mar?",
            "alternatives": ["A) 90°C", "B) 100°C", "C) 110°C", "D) 50°C", "E) 150°C"],
            "answer": "B",
            "explanation": "A água ferve a 100°C ao nível do mar.",
        },
        {
            "text": "Qual elemento é um não-metal gasoso à temperatura ambiente?",
            "alternatives": ["A) Carbono", "B) Enxofre", "C) Oxigênio", "D) Silício", "E) Fósforo"],
            "answer": "C",
            "explanation": "O oxigênio é um não-metal gasoso essencial para a respiração.",
        },
        {
            "text": "A densidade é uma propriedade física intensiva ou extensiva?",
            "alternatives": ["A) Intensiva", "B) Extensiva", "C) Ambas", "D) Nenhuma", "E) Relativa"],
            "answer": "A",
            "explanation": "A densidade é uma propriedade intensiva, pois não depende da quantidade de matéria.",
        },
        {
            "text": "Qual é o número atômico do oxigênio?",
            "alternatives": ["A) 6", "B) 7", "C) 8", "D) 9", "E) 10"],
            "answer": "C",
            "explanation": "O oxigênio tem número atômico 8.",
        },
        {
            "text": "Em uma ligação iônica, há transferência ou compartilhamento de elétrons?",
            "alternatives": ["A) Transferência", "B) Compartilhamento", "C) Fusão", "D) Perda", "E) Criação"],
            "answer": "A",
            "explanation": "Na ligação iônica há transferência de elétrons de um átomo para outro.",
        },
        {
            "text": "Qual é a fórmula química do gás oxigênio?",
            "alternatives": ["A) O", "B) O₂", "C) O₃", "D) O₄", "E) 2O"],
            "answer": "B",
            "explanation": "O oxigênio molecular é representado por O₂.",
        },
        {
            "text": "O ferro é um elemento essencial para a vida porque forma parte de qual proteína?",
            "alternatives": ["A) Proteína C", "B) Hemoglobina", "C) Colágeno", "D) Elastina", "E) Queratina"],
            "answer": "B",
            "explanation": "O ferro é essencial para a hemoglobina, que transporta oxigênio no sangue.",
        },
    ]

    q_index = 0

    for episode in episodes:
        # Episode header
        ep_header = doc.add_paragraph()
        ep_run = ep_header.add_run(f"Episódio {episode['number']} — {episode['title']}")
        ep_run.font.size = Pt(12)
        ep_run.font.bold = True
        ep_run.font.color.rgb = RGBColor(0, 102, 204)

        desc = doc.add_paragraph(episode["description"])
        desc.style = "List Bullet"

        doc.add_paragraph()

        # Add content section
        content_para = doc.add_paragraph("Conteúdo Teórico:")
        content_para.runs[0].font.bold = True

        theory_text = (
            f"Este episódio aborda os conceitos principais sobre '{episode['title'].lower()}' "
            f"e sua importância na compreensão dos princípios elementares da matéria."
        )
        doc.add_paragraph(theory_text)

        doc.add_paragraph()

        # Add 2-3 questions per episode
        questions_per_episode = 3 if episode["number"] < 5 else 2
        for _ in range(questions_per_episode):
            if q_index < len(question_bank):
                q = question_bank[q_index]

                # Question - formatted to match regex pattern "Questão N. text"
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(f"Questão {q_index + 1}. {q['text']}")
                q_run.font.bold = True

                # Alternatives - formatted as plain text (not bullet list) to match regex
                for alt in q["alternatives"]:
                    doc.add_paragraph(alt)

                # Add blank line
                doc.add_paragraph()

                q_index += 1

        doc.add_paragraph()

    # Answer Key Section
    doc.add_page_break()

    gabarito_title = doc.add_paragraph()
    gabarito_run = gabarito_title.add_run("GABARITO E COMENTÁRIOS")
    gabarito_run.font.size = Pt(14)
    gabarito_run.font.bold = True
    gabarito_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    for i, q in enumerate(question_bank):
        gabarito_para = doc.add_paragraph()
        gabarito_run = gabarito_para.add_run(f"Questão {i + 1}. Gabarito: {q['answer']}")
        gabarito_run.font.bold = True

        explicacao = doc.add_paragraph(q["explanation"], style="List Bullet")

        doc.add_paragraph()

    return doc


if __name__ == "__main__":
    doc = create_pilot_material()
    output_path = "/Users/marcoviana/agente-ia-edu-core/tests/fixtures/ingestion_materials/T 01 PRINCÍPIOS ELEMENTARES DA MATÉRIA3.docx"
    doc.save(output_path)
    print(f"✓ Pilot material created: {output_path}")

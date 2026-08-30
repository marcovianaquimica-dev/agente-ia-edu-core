"""
Deterministic DOCX/PDF parser for document ingestion.

Extracts structure, questions, sections, and metadata without using LLM.
All extraction is deterministic and repeatable.
"""

import hashlib
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.shared import OxmlElement


@dataclass
class ParsedSection:
    """A structural section extracted from a document."""

    section_type: str  # SEASON, EPISODE, CHAPTER, etc.
    title: Optional[str]
    description: Optional[str]
    section_number: Optional[str]
    page_start: Optional[int]
    page_end: Optional[int]
    content_lines: list[str]
    position: int


@dataclass
class ParsedQuestion:
    """A question extracted from a document."""

    question_number: int
    statement_text: str
    alternatives_text: Optional[str]  # newline-separated or semicolon-separated
    correct_answer: Optional[str]  # letter/number if found
    answer_explanation: Optional[str]
    page_start: Optional[int]
    page_end: Optional[int]
    position: int
    section_index: Optional[int] = None


@dataclass
class ParsedDocument:
    """Complete parsing result."""

    filename: str
    document_hash: str
    title: Optional[str]
    author: Optional[str]
    page_count: Optional[int]
    sections: list[ParsedSection]
    questions: list[ParsedQuestion]
    total_images: int = 0
    total_tables: int = 0


class DocxParser:
    """Deterministic parser for DOCX files."""

    # Regex patterns for common document structures
    SEASON_PATTERN = re.compile(r"^Temporada\s+(\d+)\s*[–:-]?\s*(.*)$", re.IGNORECASE)
    EPISODE_PATTERN = re.compile(r"^(?:Episódio|Episode|Aula|Ep\.?\s*)\s*(\d+)\s*[–:-]?\s*(.*)$", re.IGNORECASE)
    SECTION_PATTERN = re.compile(r"^(?:Seção|Section|Módulo|Bloco)\s*(\d+)?\s*[–:-]?\s*(.*)$", re.IGNORECASE)
    QUESTION_PATTERN = re.compile(r"^Questão\s+(\d+)\.\s+(.+)$", re.IGNORECASE)
    ALTERNATIVE_PATTERN = re.compile(r"^([a-eA-E])\)\s*(.+)$", re.MULTILINE)
    ANSWER_PATTERN = re.compile(r"Gabarito:\s*([a-eA-E])", re.IGNORECASE)
    EXPLANATION_PATTERN = re.compile(r"^(?:Explicação|Comentário|Explanation|Justificativa)\s*:?\s*(.+)$", re.IGNORECASE | re.MULTILINE | re.DOTALL)

    @staticmethod
    def file_hash(filepath: Path) -> str:
        """Calculate hash of file for idempotency detection."""
        sha256_hash = hashlib.sha256()
        with open(filepath, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()

    @staticmethod
    def _get_page_number(paragraph) -> Optional[int]:
        """Extract page number from paragraph if available (OOXML property)."""
        try:
            # This is a simplified approach - may need enhancement for actual .docx files
            return None
        except Exception:
            return None

    @classmethod
    def parse_file(cls, filepath: Path) -> ParsedDocument:
        """
        Parse a DOCX file deterministically.

        Returns a ParsedDocument with all extracted structure, sections, and questions.
        """
        doc = Document(filepath)

        # Extract document metadata
        core_props = doc.core_properties
        title = core_props.title or filepath.stem
        author = core_props.author

        # Extract full text and count structural elements
        sections: list[ParsedSection] = []
        questions: list[ParsedQuestion] = []
        current_section: Optional[ParsedSection] = None
        text_lines: list[str] = []
        image_count = 0
        table_count = 0
        position_counter = 0
        
        # Track if we're in answer key section (to avoid double-parsing questions)
        gabarito_started = False

        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if not para_text:
                continue

            # Check if we're entering gabarito section
            if re.match(r"^(?:GABARITO|GABARITO E COMENTÁRIOS|ANSWER KEY)", para_text, re.IGNORECASE):
                gabarito_started = True

            # Check for season header (before gabarito check)
            if not gabarito_started:
                season_match = cls.SEASON_PATTERN.match(para_text)
                if season_match:
                    if current_section:
                        sections.append(current_section)
                    current_section = ParsedSection(
                        section_type="SEASON",
                        title=f"Temporada {season_match.group(1)}",
                        description=season_match.group(2) if season_match.group(2) else None,
                        section_number=season_match.group(1),
                        page_start=None,
                        page_end=None,
                        content_lines=[],
                        position=position_counter,
                    )
                    position_counter += 1
                    text_lines.append(para_text)
                    continue

                # Check for episode header
                episode_match = cls.EPISODE_PATTERN.match(para_text)
                if episode_match:
                    if current_section:
                        sections.append(current_section)
                    current_section = ParsedSection(
                        section_type="EPISODE",
                        title=f"Episódio {episode_match.group(1)}",
                        description=episode_match.group(2) if episode_match.group(2) else None,
                        section_number=episode_match.group(1),
                        page_start=None,
                        page_end=None,
                        content_lines=[],
                        position=position_counter,
                    )
                    position_counter += 1
                    text_lines.append(para_text)
                    continue

                # Check for question (only before gabarito section)
                question_match = cls.QUESTION_PATTERN.match(para_text)
                if question_match:
                    q_number = int(question_match.group(1))
                    q_statement = question_match.group(2).strip()

                    parsed_q = ParsedQuestion(
                        question_number=q_number,
                        statement_text=q_statement,
                        alternatives_text=None,
                        correct_answer=None,
                        answer_explanation=None,
                        page_start=None,
                        page_end=None,
                        position=len(questions),
                        section_index=len(sections) - 1 if current_section else None,
                    )
                    questions.append(parsed_q)
                    text_lines.append(para_text)
                    continue

            # For all text (including gabarito section), collect it
            if current_section:
                current_section.content_lines.append(para_text)
            text_lines.append(para_text)

        # Finalize last section
        if current_section:
            sections.append(current_section)

        # Count images and tables
        image_count = len(doc.inline_shapes)
        table_count = len(doc.tables)

        # Post-process questions to extract alternatives and answers from text
        cls._enrich_questions(questions, text_lines)

        return ParsedDocument(
            filename=filepath.name,
            document_hash=cls.file_hash(filepath),
            title=title,
            author=author,
            page_count=None,  # DOCX doesn't easily expose page count without python-pptx
            sections=sections,
            questions=questions,
            total_images=image_count,
            total_tables=table_count,
        )

    @classmethod
    def _enrich_questions(cls, questions: list[ParsedQuestion], text_lines: list[str]) -> None:
        """
        Post-process questions to extract alternatives and answer key.

        Looks for patterns like "a) Alternative text", "Gabarito: A", etc.
        """
        # Join all text lines for searching
        full_text = "\n".join(text_lines)

        for question in questions:
            # Find this question's statement in the full text
            q_statement = question.statement_text
            
            # Try to find the statement
            statement_idx = None
            for i, line in enumerate(text_lines):
                if q_statement in line or line in q_statement:
                    statement_idx = i
                    break
            
            if statement_idx is not None:
                # Extract text from after this statement until next question or end
                # Look ahead for alternatives
                alternatives = []
                for i in range(statement_idx + 1, min(statement_idx + 20, len(text_lines))):
                    line = text_lines[i].strip()
                    
                    # Stop if we hit another question
                    if line.startswith("Questão"):
                        break
                    
                    # Check for alternative pattern: "A) text" or "a) text"
                    alt_match = re.match(r"^([a-eA-E])\)\s*(.+)$", line)
                    if alt_match:
                        letter = alt_match.group(1).upper()
                        text = alt_match.group(2).strip()
                        alternatives.append(f"{letter}) {text}")
                    elif line and alternatives and not alt_match:
                        # Stop collecting alternatives if we hit non-alternative text
                        break
                
                if alternatives:
                    question.alternatives_text = "\n".join(alternatives)
            
            # Look for answer key in gabarito section (separate search in full_text)
            # Pattern: "Questão N. Gabarito: X"
            gabarito_pattern = f"Questão {question.question_number}.*?Gabarito:\\s*([A-Ea-e])"
            gabarito_match = re.search(gabarito_pattern, full_text, re.IGNORECASE | re.DOTALL)
            if gabarito_match:
                question.correct_answer = gabarito_match.group(1).upper()


class PdfParser:
    """
    Placeholder for PDF parsing (not yet implemented in MVP).

    PDF extraction would require additional dependencies (PyPDF2, pdfplumber, etc.)
    and more complex logic to handle variable formatting.
    """

    @staticmethod
    def parse_file(filepath: Path) -> ParsedDocument:
        raise NotImplementedError("PDF parsing will be implemented in a future phase")


def parse_document(filepath: Path) -> ParsedDocument:
    """
    Automatic dispatch to the appropriate parser based on file extension.

    Deterministic, no LLM, repeatable for identical inputs.
    """
    suffix = filepath.suffix.lower()

    if suffix == ".docx":
        return DocxParser.parse_file(filepath)
    elif suffix == ".pdf":
        return PdfParser.parse_file(filepath)
    else:
        raise ValueError(f"Unsupported document format: {suffix}")
